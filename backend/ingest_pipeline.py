"""
Pipeline de ingestão de documentos.
Lida com a extração de texto (PDF, DOCX, TXT), fragmentação (chunking),
geração de embeddings e indexação no Qdrant com metadados enriquecidos.
"""

import os
import hashlib
import json
import uuid
import logging
from datetime import datetime, timezone
from typing import Optional, List, Dict

logger = logging.getLogger(__name__)

import fitz
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from qdrant_client import QdrantClient
from qdrant_client.http import models
import ollama

from backend.config import settings
from backend.auth import validate_department
from backend.graph_engine import extract_graph_data, add_graph_data


_qdrant_client: Optional[QdrantClient] = None

def get_qdrant_client() -> QdrantClient:
    """Retorna ou cria o cliente Qdrant."""
    global _qdrant_client
    if _qdrant_client is None:
        _qdrant_client = QdrantClient(
            host=settings.qdrant_host,
            port=settings.qdrant_port
        )
        _ensure_collection_exists(_qdrant_client)
    return _qdrant_client


def _ensure_collection_exists(client: QdrantClient):
    """Garante que a coleção exista no Qdrant com a configuração correta."""
    try:
        collections = client.get_collections().collections
        exists = any(c.name == settings.qdrant_collection_name for c in collections)
        
        if not exists:
            logger.info(f"Criando coleção '{settings.qdrant_collection_name}'...")
            client.create_collection(
                collection_name=settings.qdrant_collection_name,
                vectors_config=models.VectorParams(
                    size=768, 
                    distance=models.Distance.COSINE
                ),
            )
            logger.info("Coleção criada com sucesso.")
    except Exception as e:
        logger.error(f"Erro ao verificar/criar coleção: {e}")


def extract_text_from_pdf(file_path: str) -> List[Dict]:
    """Extrai texto de PDF usando PyMuPDF."""
    pages = []
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        if text.strip():
            section = _detect_section_header(page)
            pages.append({
                "text": text.strip(),
                "page": page_num + 1,
                "section": section,
            })
    doc.close()
    return pages


def _detect_section_header(page) -> str:
    """Tenta detectar um cabeçalho de seção na página."""
    blocks = page.get_text("dict")["blocks"]
    for block in blocks:
        if "lines" in block:
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    size = span["size"]
                    if text and size >= 14:
                        return text[:100]
    return ""


def extract_text_from_docx(file_path: str) -> List[Dict]:
    """Extrai texto de DOCX (parágrafos e tabelas)."""
    doc = DocxDocument(file_path)
    pages = []
    current_text = []
    

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current_text.append(text)
            

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                current_text.append(" | ".join(row_text))

    if not current_text:
        return []


    full_text = "\n".join(current_text)
    return [{"text": full_text, "page": 1, "section": "Documento Completo"}]


def extract_text_from_txt(file_path: str) -> List[Dict]:
    """Extrai texto de arquivo TXT."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    if not text.strip():
        return []
    return [{"text": text.strip(), "page": 1, "section": ""}]


def extract_text(file_path: str, file_type: str) -> List[Dict]:
    """Roteia a extração com base no tipo de arquivo."""
    extractors = {"pdf": extract_text_from_pdf, "docx": extract_text_from_docx, "txt": extract_text_from_txt}
    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Tipo de arquivo não suportado: {file_type}")
    return extractor(file_path)


def chunk_text(pages: List[Dict]) -> List[Dict]:
    """Divide o texto em fragmentos (chunks)."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )

    chunks = []
    chunk_index = 0
    for page_data in pages:
        split_texts = splitter.split_text(page_data["text"])
        for split_text in split_texts:
            chunks.append({
                "text": split_text,
                "page": page_data["page"],
                "section": page_data["section"],
                "chunk_index": chunk_index,
            })
            chunk_index += 1
    return chunks


def generate_embeddings_batch(texts: List[str]) -> List[List[float]]:
    """Gera embeddings em lote usando Ollama."""
    response = ollama.embed(model=settings.embedding_model, input=texts)
    return response["embeddings"]


def ingest_document(file_path: str, document_name: str, department: str = "public", file_type: str = "pdf") -> Dict:
    """
    Pipeline completo: Extração -> Fragmentação -> Embeddings -> Qdrant.
    """
    client = get_qdrant_client()
    department = validate_department(department)
    timestamp = datetime.now(timezone.utc).isoformat()


    pages = extract_text(file_path, file_type)
    if not pages:
        return {"document_id": "", "document_name": document_name, "chunk_count": 0, "status": "error", "message": "Nenhum texto extraído."}

    total_pages = max(p["page"] for p in pages)


    chunks = chunk_text(pages)
    if not chunks:
        return {"document_id": "", "document_name": document_name, "chunk_count": 0, "status": "error", "message": "Falha na fragmentação."}


    version = _get_next_version(document_name)
    if version > 1:
        delete_document(document_name)

    doc_id = hashlib.sha256(document_name.encode()).hexdigest()[:16] + f"_v{version}"


    trust_score = _calculate_trust_score(document_name, file_type)


    chunk_texts = [c["text"] for c in chunks]
    all_embeddings = []
    batch_size = 50
    for i in range(0, len(chunk_texts), batch_size):
        batch = chunk_texts[i:i + batch_size]
        all_embeddings.extend(generate_embeddings_batch(batch))


    points = []
    for i, (text, vector, chunk) in enumerate(zip(chunk_texts, all_embeddings, chunks)):
        point_id = str(uuid.uuid4())
        payload = {
            "text": text,
            "document_name": document_name,
            "document_id": doc_id,
            "page": chunk["page"],
            "section": chunk["section"],
            "chunk_index": chunk["chunk_index"],
            "department": department,
            "timestamp": timestamp,
            "file_type": file_type,
            "version": version,
            "trust_score": trust_score,
        }
        points.append(models.PointStruct(id=point_id, vector=vector, payload=payload))


        if (chunk["chunk_index"] == 0 or (i > 0 and chunk["page"] != chunks[i-1]["page"])) and chunk["page"] <= 3:
             graph_data = extract_graph_data(text)
             add_graph_data(document_name, graph_data)

    client.upsert(
        collection_name=settings.qdrant_collection_name,
        points=points
    )


    _save_document_registry(doc_id, document_name, department, file_type, total_pages, len(chunks), timestamp, version)

    return {
        "document_id": doc_id,
        "document_name": document_name,
        "chunk_count": len(chunks),
        "status": "success",
        "message": f"Sucesso: {len(chunks)} fragmentos ingeridos (v{version})."
    }


def _get_next_version(document_name: str) -> int:
    """Verifica versão atual no registro e retorna a próxima."""
    registry = get_indexed_documents_dict()
    if document_name in registry:
        return registry[document_name].get("version", 1) + 1
    return 1


def delete_document(document_name: str) -> bool:
    """Remove um documento do Qdrant e do registro."""
    client = get_qdrant_client()

    client.delete(
        collection_name=settings.qdrant_collection_name,
        points_selector=models.FilterSelector(
            filter=models.Filter(
                must=[
                    models.FieldCondition(
                        key="document_name",
                        match=models.MatchValue(value=document_name)
                    )
                ]
            )
        )
    )
    _remove_from_registry(document_name)
    return True


def get_indexed_documents() -> List[Dict]:
    """Retorna lista de documentos indexados."""
    return list(get_indexed_documents_dict().values())


def get_indexed_documents_dict() -> Dict:
    """Carrega o dicionário do registro."""
    if not os.path.exists(settings.metadata_registry):
        return {}
    with open(settings.metadata_registry, "r", encoding="utf-8") as f:
        try:
            return json.load(f)
        except:
            return {}


def _save_document_registry(doc_id, name, department, file_type, total_pages, total_chunks, timestamp, version):
    """Salva metadados no arquivo JSON de registro."""
    registry = get_indexed_documents_dict()
    registry[name] = {
        "document_id": doc_id,
        "document_name": name,
        "department": department,
        "file_type": file_type,
        "total_pages": total_pages,
        "total_chunks": total_chunks,
        "upload_timestamp": timestamp,
        "version": version,
    }
    os.makedirs(os.path.dirname(settings.metadata_registry), exist_ok=True)
    with open(settings.metadata_registry, "w", encoding="utf-8") as f:
        json.dump(registry, f, indent=2, ensure_ascii=False)


def _remove_from_registry(document_name: str):
    """Remove do registro JSON."""
    registry = get_indexed_documents_dict()
    if document_name in registry:
        registry.pop(document_name)
        with open(settings.metadata_registry, "w", encoding="utf-8") as f:
            json.dump(registry, f, indent=2, ensure_ascii=False)


def get_collection():
    """Helper para manter compatibilidade com API health check (agora verifica Qdrant)."""
    return get_qdrant_client().get_collection(settings.qdrant_collection_name)


def _calculate_trust_score(document_name: str, file_type: str) -> float:
    """
    Calcula o score de confiabilidade de um documento baseado em:
    - Tipo de arquivo (PDF = alto, TXT = baixo)
    - Padrões no nome do arquivo (policy, manual = alto; rascunho, draft = baixo)
    Retorna um valor entre 0.0 e 1.0.
    """

    type_score = settings.trust_type_scores.get(file_type.lower(), 0.5)


    name_lower = document_name.lower()
    name_score = 0.7
    for pattern, score in settings.trust_name_patterns.items():
        if pattern in name_lower:
            name_score = max(name_score, score)


    final_score = (0.4 * type_score) + (0.6 * name_score)
    return round(min(1.0, max(0.0, final_score)), 2)
